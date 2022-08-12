from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
import json
import numpy as np
import itertools
import jsonlines
import os 

files = os.listdir('data/NaCha/inference_data_out')
for file in files:
    if file.endswith('jsonl'):
        print('data/NaCha/word/'+file.split('.jsonl')[0]+'.docx')
        data = []
        with jsonlines.open('data/NaCha/inference_data_out/'+file) as f:
            for line in f.iter():
                data.append(line)

        #get all the different entity categories present in the data
        entity_types = []
        for elem in data:
            if len(elem['predicted_ner'][0])>0:
                for mention in elem['predicted_ner'][0]:
                    entity_types.append(mention[2])
        unique_entity_types = np.unique(entity_types)

        data_indices = []
        data_types = []
        for elem in data:
            indices = []
            types = {}
            for type in unique_entity_types:
                types[type] = []
            if len(elem['predicted_ner'][0]) >0:
                for mention in elem['predicted_ner'][0]:
                    indices.append(list(np.arange(mention[0], mention[1]+1)))
                    types[mention[2]].extend(list(np.arange(mention[0], mention[1]+1)))
            else:
                indices.append([])
            data_indices.append(list(itertools.chain(*indices)))
            data_types.append(types)
        #print(data_indices)
        #print(data_types)

        document = Document()
        p = document.add_paragraph('')

        colors = []
        colors.append(RGBColor(204,204,0))
        colors.append(RGBColor(102, 204, 0))
        colors.append(RGBColor(255, 0, 127))
        colors.append(RGBColor(0, 0, 255))
        colors.append(RGBColor(255, 51, 51))

        for j in range(len(unique_entity_types)):
            p.add_run().add_break()
            p.add_run(unique_entity_types[j]).font.color.rgb = colors[j]

        for k in range(len(data)):
            p.add_run().add_break()
            p.add_run().add_break()
            for i in range(len(data[k]['sentences'][0])):
                condition = False
                for f in range(len(unique_entity_types)):
                    if i in data_types[k][unique_entity_types[f]]:
                        p.add_run(data[k]['sentences'][0][i]+' ').font.color.rgb = colors[f]
                        condition = True
                if condition == False:
                    p.add_run(data[k]['sentences'][0][i]+' ').font.color.rgb = RGBColor(0,0,0)
            


        document.add_page_break()
        document.save('data/NaCha/word/'+file.split('.jsonl')[0]+'.docx')