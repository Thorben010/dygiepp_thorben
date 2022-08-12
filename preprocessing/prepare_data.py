import json
import pandas as pd
from nltk.tokenize import word_tokenize
import jsonlines
import numpy as np
import pickle5 as pickle

class preprocessor:
    def __init__(self, path = 'data/challenges_papers_with_mat_category.pickle'):
        #with open(path, 'r') as f:
        #    self.data = json.load(f)

        with open(path, 'rb') as handle:
            self.data = pickle.load(handle)
        return


    def create_jsonl(self):
        sections = ['abstract', 'intro', 'results', 'conclusions']
        list_dicts = []
        for i in range(len(self.data)):
            paper_dict = {}
            paper_dict['dataset'] = "covid-event"
            paper_dict['doc_key'] = self.data[i]['title']
            paper_dict['events'] = []

            list_sentences = []
            for section in sections:
                if section in self.data[i]:
                    for sentence in self.data[i][section]['challenge']:
                        list_sentences.append(word_tokenize(sentence))
                        paper_dict['events'].append([])
                paper_dict['sentences'] = list_sentences
                if len(paper_dict['sentences']) >0:
                    list_dicts.append(paper_dict)

            with jsonlines.open('data/test_dict.jsonl', mode='w') as writer:
                for elem in list_dicts:
                    writer.write(elem)
        print('done')

#obj = preprocessor()
#obj.data = obj.data['MnFe']
#obj.create_jsonl()

"""
allennlp predict \
  pretrained/mechanic-coarse.tar.gz \
  data/5-6_important_challenges_test.json \
  --predictor dygie \
  --include-package dygie \
  --use-dataset-reader \
  --output-file predictions/5-6_important_challenges_test.jsonl \
  --cuda-device -1 \
  --silent
"""


#dict = {"doc_key": "High Reversible Pseudocapacity in Mesoporous Yolk\u2013Shell Anatase TiO2/TiO2(B) Microspheres Used as Anodes for Li\u2010Ion Batteries", "dataset": "covid-event", "sentences": [["To", "address", "the", "poor", "rate", "behavior", ",", "slow", "lithium-ion", "(", "Li+", ")", "diffusion", ",", "and", "high", "irreversible", "capacity", "decay", ",", "TiO2", "nanomaterials", "with", "tuned", "phase", "compositions", "and", "morphologies", "are", "being", "investigated", "."], ["However", ",", "the", "practical", "application", "of", "anatase", "TiO2", "has", "been", "hampered", "by", "its", "low", "rate", "performance", "due", "to", "high", "resistance", "and", "poor", "lithium-ion", "(", "Li+", ")", "diffusivity", "at", "the", "electrode/electrolyte", "interface", "."], ["To", "tackle", "these", "issues", ",", "numerous", "strategies", "have", "been", "tested", ",", "including", "decreasing", "the", "size", "of", "the", "particles", "to", "the", "nanoscale", ",", "increasing", "porosity", ",", "and", "the", "addition", "of", "other", "conductive", "materials", "such", "as", "carbon", ",", "SnO2", ",", "and", "V2O5", "."], ["However", ",", "TiO2", "(", "B", ")", "also", "suffers", "some", "intrinsic", "disadvantages", ",", "such", "as", "poor", "electronic", "conductivity", "and", "Li+", "diffusivity", "."], ["While", "reducing", "the", "particle", "size", "alleviates", "these", "issues", ",", "it", "has", "been", "found", "that", "nanostructured", "TiO2", "(", "B", ")", "exhibits", "significant", "irreversible", "capacity", "decay", "upon", "the", "first", "discharge", "due", "to", "reactive", "TiOH", "and", "TiO", "surface", "sites", "that", ",", "it", "is", "suggested", ",", "could", "cause", "electrolyte", "degradation", "and", "irreversible", "Li+", "trapping", "."], ["Due", "to", "the", "thermal", "instability", "of", "the", "metastable", "crystal", "structure", ",", "the", "traces", "of", "the", "TiO2", "(", "B", ")", "phase", "(", "<", "5", "wt", "%", ")", "that", "remained", "at", "high", "temperature", "(", ">", "500", "degC", ")", "form", "limited", "anatase/TiO2", "(", "B", ")", "interfaces", "."], ["Monodisperse", "microspheres", "were", "not", "obtained", "by", "applying", "the", "same", "procedure", "to", "raw", "materials", "of", "smaller", "particle", "size", "(", "300", "nm", ")", "."], ["In", "contrast", ",", "the", "corresponding", "charge", "transfer", "resistance", "of", "MM", "(", "268", "O", ")", "was", "approximately", "fivefold", "higher", "than", "that", "of", "YSM", ",", "indicating", "that", "low", "surface", "area", "and", "limited", "internal", "pore", "volume", "could", "hamper", "mass", "and", "charge", "transfer", "."], ["MM", "and", "CSM", "revealed", "capacity", "decay", "over", "the", "500", "cycles", ",", "with", "79", "%", "and", "89", "%", "retention", ",", "respectively", "."], ["In", "this", "case", "though", ",", "the", "cyclability", "is", "relatively", "poor", "when", "compared", "with", "YSM", ",", "and", "as", "the", "discharge", "rate", "was", "lowered", ",", "the", "capacity", "of", "black", "anatase", "TiO2", "dropped", "significantly", "."]], "events": [[], [], [], [], [], [], [], [], [], []], "predicted_events": [[], [[[48, "TRIGGER", 15.9299, 1.0], [38, 39, "ARG0", 3.9275, 0.7783], [45, 47, "ARG1", 2.5832, 0.9294], [50, 51, "ARG1", 3.6109, 0.8126], [53, 62, "ARG1", 0.2551, 0.5633]]], [[[76, "TRIGGER", 13.3178, 1.0], [78, 84, "ARG1", 8.8211, 0.9999]], [[86, "TRIGGER", 9.833, 0.9999], [78, 84, "ARG1", 1.3382, 0.7915], [87, 87, "ARG1", 4.2751, 0.9166]]], [[[112, "TRIGGER", 9.0234, 0.9999], [107, 110, "ARG0", 7.8551, 0.9979], [119, 121, "ARG1", 5.746, 0.9968], [120, 121, "ARG1", 0.8725, 0.7053], [123, 124, "ARG1", 5.2108, 0.9945]]], [[[145, "TRIGGER", 1.0334, 0.7376], [140, 144, "ARG0", 1.0119, 0.733]], [[154, "TRIGGER", 17.9625, 1.0], [147, 153, "ARG1", 2.2231, 0.9021], [156, 161, "ARG0", 4.8943, 0.9921]], [[169, "TRIGGER", 16.0157, 1.0], [170, 171, "ARG1", 9.0878, 0.9999], [173, 175, "ARG1", 7.1047, 0.9992]]], [[[177, "TRIGGER", 2.4099, 0.9176], [180, 186, "ARG0", 4.137, 0.9817]], [[213, "TRIGGER", 11.5553, 1.0]]], [], [[[277, "TRIGGER", 14.5984, 1.0], [268, 275, "ARG0", 8.9685, 0.9999], [272, 275, "ARG0", 3.0093, 0.953], [278, 281, "ARG1", 11.7495, 1.0]]], [[[286, "TRIGGER", 17.5341, 1.0], [285, 285, "ARG0", 16.3915, 1.0], [287, 288, "ARG1", 9.1942, 0.9999]]], []]}



def interpret_jasonl_lines(df, dict):
    for i in range(len(dict['sentences'])):
        if i != 0:
            length_sentence = 0
            for kl in range(i):
                length_sentence += len(dict['sentences'][kl])
        else:
            length_sentence = 0

        if len(dict['predicted_events'][i]) > 0:
            for k in range(len(dict['predicted_events'][i][0])):
                dict_df = {}
                if len(dict['predicted_events'][i][0][k]) == 5:
                    if dict['predicted_events'][i][0][k][4] > 0.9:
                        dict_df['sentence'] = dict['sentences'][i]
                        start = int(dict['predicted_events'][i][0][k][0])- length_sentence
                        end = int(dict['predicted_events'][i][0][k][1])- length_sentence
                        dict_df['start'] = start
                        dict_df['end'] = end
                        dict_df['type'] = dict['predicted_events'][i][0][k][2]
                        dict_df['probability'] = dict['predicted_events'][i][0][k][4]
                        if start == end:
                            dict_df['tokens'] = dict['sentences'][i][start]
                        else:
                            dict_df['tokens'] = dict['sentences'][i][start:end+1]
                if len(dict['predicted_events'][i][0][k]) == 4:
                    if dict['predicted_events'][i][0][k][3] > 0.9:
                        dict_df['sentence'] = dict['sentences'][i]
                        dict_df['start'] = dict['predicted_events'][i][0][k][0] - length_sentence
                        dict_df['end'] = None
                        dict_df['type'] = dict['predicted_events'][i][0][k][1]
                        dict_df['probability'] = dict['predicted_events'][i][0][k][3]
                        dict_df['tokens'] = dict['sentences'][i][dict['predicted_events'][i][0][k][0] - length_sentence]
                if len(dict_df) > 0:
                    df = df.append(dict_df, ignore_index=True)
    return df

df = pd.DataFrame()
with jsonlines.open('predictions/covid-coarse_1.jsonl') as f:
    for line in f.iter():
        df = interpret_jasonl_lines(df, line)


def extract_triplets(df):
    numbers = []
    for i in range(len(df)):
        if i != 0:
            if df.loc[i]['sentence'] == df.loc[i-1]['sentence']:
                if df.loc[i]['type'] == 'TRIGGER':
                    check[0] = 1
                if df.loc[i]['type'] == 'ARG0':
                    check[1] = 1
                if df.loc[i]['type'] == 'ARG1':
                    check[2] = 1
                indices.append(i)
            else:
                if sum(check) == 3:
                    numbers.append(indices)
                check = [0, 0, 0]
                indices = []
                if df.loc[i]['type'] == 'TRIGGER':
                    check[0] = 1
                if df.loc[i]['type'] == 'ARG0':
                    check[1] = 1
                if df.loc[i]['type'] == 'ARG1':
                    check[2] = 1
                indices.append(i)
        else:
            check = [0, 0, 0]
            indices = []
    return df.loc[np.array([element for sublist in numbers for element in sublist])]


new_df = extract_triplets(df)


results = []
for i in range(len(new_df)):
    if new_df.iloc[i]['end'] != None:
        if new_df.iloc[i]['end'] != new_df.iloc[i]['start']:
            results.append(list(range(int(new_df.iloc[i]['start']), int(new_df.iloc[i]['end']+1))))
        else:
            results.append([int(new_df.iloc[i]['start'])])
    else:
        results.append([int(new_df.iloc[i]['start'])])
new_df["index"] = results


def write_word_out(new_df):
    new_df['sentence'] = new_df['sentence'].apply(lambda x: ' '.join(x))
    uniques = new_df['sentence'].unique()
    dict = {}
    for item in uniques:
        dict[item] = {}
        indices = list(new_df.iloc[np.where((new_df['type'] == 'ARG0') & (new_df['sentence'] == item))]['index'])
        if type(indices) == list and len(indices) > 1:
            dict[item]['ARG0'] = [element for sublist in indices for element in sublist]
        else:
            dict[item]['ARG0'] = indices[0]
        indices = list(new_df.iloc[np.where((new_df['type'] == 'ARG1') & (new_df['sentence'] == item))]['index'])
        if type(indices) == list and len(indices) > 1:
            dict[item]['ARG1'] = [element for sublist in indices for element in sublist]
        else:
            dict[item]['ARG1'] = indices[0]
        indices = list(new_df.iloc[np.where((new_df['type'] == 'TRIGGER') & (new_df['sentence'] == item))]['index'])
        if len(indices) > 1:
            dict[item]['TRIGGER'] = [element for sublist in indices for element in sublist]
        else:
            dict[item]['TRIGGER'] = indices[0]
    return dict

dict = write_word_out(new_df)

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

document = Document()
p = document.add_paragraph('Hi')

for sentence in dict:
    p.add_run().add_break()
    p.add_run().add_break()
    position = 0
    for word in word_tokenize(sentence):
        if position in dict[sentence]['ARG1']:
            p.add_run(word+' ').font.color.rgb = RGBColor(255, 20, 20)
        else:
            if position in dict[sentence]['ARG0']:
                p.add_run(word+' ').font.color.rgb = RGBColor(20, 20, 255)
            else:
                if position in dict[sentence]['TRIGGER']:
                    p.add_run(word+' ').bold = True
                else:
                    p.add_run(word+' ')
        position += 1

document.add_page_break()
document.save('extraction_results_both_NaMnO.docx')
print()